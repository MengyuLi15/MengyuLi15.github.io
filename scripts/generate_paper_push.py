from __future__ import annotations

import argparse
import concurrent.futures
import csv
import html
import json
import os
import re
import subprocess
import sys
import textwrap
import time
import unicodedata
import urllib.error
import urllib.parse
import urllib.request
from dataclasses import dataclass
from datetime import date, datetime, timedelta, timezone
from functools import lru_cache
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "_data" / "paper_pushes.yml"
HISTORY_PATH = ROOT / "推送历史_论文索引.csv"
DATE_PAGE_SCRIPT = ROOT / "scripts" / "create_paper_push_page.py"
FOCUSED_TEAM_AUTHORS_PATH = ROOT / "scripts" / "focused_team_authors.json"
TZ = ZoneInfo("Europe/Berlin")

MAX_ABSTRACT_CHARS = 900
USER_AGENT = "mengyuli15-paper-push/1.0 (https://mengyuli15.github.io/)"

TOPIC_TERMS = [
    "BGC Argo",
    "BGC-Argo",
    "biogeochemical Argo",
    "BGC-Argo carbon pump",
    "BGC-Argo phytoplankton",
    "BGC-Argo microbial carbon",
    "BGC-Argo bio-optic",
    "biological carbon pump ocean",
    "ocean carbon export",
    "particulate organic carbon ocean",
    "net community production ocean",
    "marine phytoplankton",
    "ocean phytoplankton",
    "phytoplankton community marine",
    "marine chlorophyll remote sensing",
    "ocean colour remote sensing phytoplankton",
    "ocean color remote sensing phytoplankton",
    "ocean bio-optics",
    "bio-optical ocean",
    "ocean optics absorption backscattering",
    "ocean optical absorption",
    "marine inherent optical properties",
    "in situ inherent optical properties ocean",
    "in situ apparent optical properties ocean",
    "in situ IOP AOP ocean optics",
    "seawater optical properties",
    "seawater angular scattering function",
    "seawater backscattering particles",
    "ocean IOP absorption backscattering",
    "ocean AOP remote sensing reflectance",
    "apparent optical properties ocean color",
    "particulate backscattering ocean",
    "particle backscattering ocean",
    "very small particles seawater backscattering",
    "phytoplankton absorption ocean",
    "particle absorption coefficient ocean",
    "particulate absorption coefficient ocean",
    "particulate absorption coefficients hyperspectral",
    "hyperspectral particulate absorption coefficients",
    "IEEE TGRS hyperspectral particulate absorption coefficients",
    "mesoscale eddies particulate absorption coefficients",
    "hyperspectral ocean optics absorption",
    "colored dissolved organic matter absorption ocean",
    "CDOM absorption ocean",
    "absorption coefficient ocean color",
    "backscattering coefficient ocean color",
    "remote sensing reflectance IOP ocean",
    "bio-optical algorithms ocean absorption",
    "aquatic optics ocean color",
    "ocean optics applications",
    "marine bacterioplankton carbon",
    "microbial carbon ocean",
    "marine heatwave phytoplankton",
    "marine heatwave chlorophyll",
    "marine heatwave phytoplankton vertical",
    "marine heatwave deep chlorophyll maximum",
    "vertical phytoplankton ocean",
    "deep chlorophyll maximum ocean",
    "particulate organic carbon ocean bio-optics",
    "particulate organic carbon satellites Argo floats",
    "satellite ship Argo particulate organic carbon",
    "consistent estimates particulate organic carbon satellites ships Argo floats",
    "surface ocean chlorophyll trend",
    "surface ocean chlorophyll climate",
    "surface ocean chlorophyll green get greener blue bluer",
    "ocean chlorophyll warming trend",
    "long term trends mesoscale eddies hyperspectral particulate absorption coefficients",
    "backscattering particles sizes 0.2 seawater",
]

NATURE_SCIENCE_HINTS = [
    "Nature",
    "Nature Climate Change",
    "Nature Communications",
    "Communications Earth & Environment",
    "Science",
    "Science Advances",
]

PRIORITY_JOURNALS = [
    "AGU Advances",
    "Annual Review of Marine Science",
    "Applied Optics",
    "Biogeosciences",
    "Communications Earth & Environment",
    "Deep Sea Research Part I: Oceanographic Research Papers",
    "Deep Sea Research Part II: Topical Studies in Oceanography",
    "Earth System Science Data",
    "Environmental Research Letters",
    "Frontiers in Marine Science",
    "Geophysical Research Letters",
    "Global Biogeochemical Cycles",
    "Global Change Biology",
    "IEEE Transactions on Geoscience and Remote Sensing",
    "Journal of Atmospheric and Oceanic Technology",
    "Journal of Geophysical Research: Oceans",
    "Journal of Quantitative Spectroscopy and Radiative Transfer",
    "Limnology and Oceanography",
    "Limnology and Oceanography: Methods",
    "Oceanography",
    "Optics Express",
    "Proceedings of the National Academy of Sciences",
    "Progress in Oceanography",
    "Remote Sensing of Environment",
    "Remote Sensing",
]

# Approximate ordering by recent impact/influence and field relevance.
JOURNAL_RANK = {
    "Proceedings of the National Academy of Sciences": 10,
    "Remote Sensing of Environment": 20,
    "Global Change Biology": 30,
    "Annual Review of Marine Science": 40,
    "AGU Advances": 50,
    "Communications Earth & Environment": 60,
    "Geophysical Research Letters": 70,
    "Global Biogeochemical Cycles": 80,
    "Earth System Science Data": 90,
    "Environmental Research Letters": 100,
    "Journal of Geophysical Research: Oceans": 110,
    "Limnology and Oceanography": 120,
    "Progress in Oceanography": 130,
    "Biogeosciences": 140,
    "IEEE Transactions on Geoscience and Remote Sensing": 150,
    "Frontiers in Marine Science": 160,
    "Deep Sea Research Part I: Oceanographic Research Papers": 170,
    "Deep Sea Research Part II: Topical Studies in Oceanography": 180,
    "Journal of Atmospheric and Oceanic Technology": 190,
    "Limnology and Oceanography: Methods": 200,
    "Oceanography": 210,
    "Optics Express": 220,
    "Applied Optics": 230,
    "Journal of Quantitative Spectroscopy and Radiative Transfer": 240,
    "Remote Sensing": 250,
}

RELEVANCE_TERMS = {
    "bgc-argo": 8,
    "biogeochemical-argo": 8,
    "biogeochemical argo": 8,
    "argo float": 4,
    "carbon pump": 7,
    "carbon export": 6,
    "net community production": 6,
    "particulate organic carbon": 5,
    "poc": 4,
    "microbial carbon": 6,
    "bacterioplankton": 4,
    "microbiome": 4,
    "phytoplankton": 5,
    "chlorophyll": 4,
    "deep chlorophyll maximum": 6,
    "subsurface chlorophyll": 6,
    "primary production": 5,
    "net primary production": 6,
    "ocean colour": 5,
    "ocean color": 5,
    "ocean optics": 6,
    "aquatic optics": 5,
    "bio-optic": 5,
    "bio-optical": 5,
    "inherent optical properties": 6,
    "apparent optical properties": 5,
    "apparent optical property": 5,
    "iop": 4,
    "iops": 4,
    "aop": 4,
    "aops": 4,
    "in situ": 2,
    "in-situ": 2,
    "optical absorption": 6,
    "light absorption": 4,
    "absorption coefficient": 6,
    "absorption coefficients": 6,
    "phytoplankton absorption": 6,
    "particulate absorption": 5,
    "particulate absorption coefficient": 7,
    "particulate absorption coefficients": 7,
    "particle absorption": 5,
    "cdom absorption": 5,
    "colored dissolved organic matter": 5,
    "coloured dissolved organic matter": 5,
    "backscatter": 4,
    "backscattering": 5,
    "backscattering coefficient": 6,
    "particulate backscattering": 6,
    "particle backscattering": 5,
    "angular scattering": 6,
    "volume scattering": 5,
    "very small particles": 5,
    "bbp": 5,
    "aph": 4,
    "adg": 4,
    "vsp": 4,
    "remote sensing reflectance": 5,
    "rrs": 4,
    "diffuse attenuation": 5,
    "kd": 4,
    "water-leaving radiance": 4,
    "ocean lidar": 5,
    "remote sensing": 4,
    "satellite": 3,
    "hyperspectral": 4,
    "surface ocean chlorophyll": 6,
    "marine heatwave": 7,
    "marine heatwaves": 7,
    "temperature sensitivity": 3,
    "vertical": 2,
}

EXCLUDED_TITLE_PREFIXES = ("corrigendum", "erratum", "correction", "retraction")
DANTE_CARDS = [
    {
        "phrase": "Nel mezzo del cammin di nostra vita, mi ritrovai per una selva oscura.",
        "source": "Dante, Commedia, Inferno I, 1-2; Italian original from Kalliope",
        "explanation_zh": "这是 Kalliope 所列《神曲》意大利语原文。现代意大利语可理解为“在人生旅程的中途，我发现自己走入一片黑暗森林”；常用来表达迷茫、转折和重新寻找方向。",
        "explanation_en": "This follows the Italian original listed by Kalliope. In modern terms, it means that midway through life's journey, the speaker finds himself in a dark wood, a scene of disorientation and renewed searching.",
    },
    {
        "phrase": "Lasciate ogne speranza, voi ch'intrate.",
        "source": "Dante, Commedia, Inferno III, 9; Italian original from Kalliope",
        "explanation_zh": "这是地狱门上的警句，意为“进入这里的人，放弃一切希望”。原文里的 ogne 和 ch'intrate 保留了中古意大利语色彩。",
        "explanation_en": "This is the warning at the gate of Hell: those who enter must abandon hope. Today it can describe a severe situation with almost no easy way back.",
    },
    {
        "phrase": "L'amor che move il sole e l'altre stelle.",
        "source": "Dante, Commedia, Paradiso XXXIII, 145; Italian original from Kalliope tradition",
        "explanation_zh": "《神曲》的结尾句，把“爱”写成推动太阳和群星运行的力量；原文 amor、move、l'altre 带有古典/中古意大利语形式。",
        "explanation_en": "The final line of the Divine Comedy presents love as the force moving the sun and the stars, linking love with order and cosmic meaning.",
    },
    {
        "phrase": "Fatti non foste a viver come bruti, ma per seguir virtute e canoscenza.",
        "source": "Dante, Commedia, Inferno XXVI, 119-120; Italian original from Kalliope",
        "explanation_zh": "这句话说人不应像野兽一样活着，而应追求德性与知识；virtute 和 canoscenza 是原文中古意大利语形式。",
        "explanation_en": "The line says humans were not made to live like beasts, but to pursue virtue and knowledge; it is often read as a call to learning and exploration.",
    },
    {
        "phrase": "E quindi uscimmo a riveder le stelle.",
        "source": "Dante, Commedia, Inferno XXXIV, 139; Italian original from Kalliope",
        "explanation_zh": "这是《地狱篇》的最后一句，意思是“于是我们走出那里，再次看见群星”。它常被用来表达穿过黑暗之后重新看见方向和希望。",
        "explanation_en": "This is the final line of Inferno: after passing through darkness, the travellers come out to see the stars again. It suggests recovery, orientation, and renewed hope.",
    },
    {
        "phrase": "Amor, ch'a nullo amato amar perdona.",
        "source": "Dante, Commedia, Inferno V, 103; Italian original from Kalliope",
        "explanation_zh": "这句来自 Francesca 的叙述，大意是“爱使被爱者不能不回以爱”。原文的 amor 和 amar 保留了诗歌中的紧密回环。",
        "explanation_en": "In Francesca's speech, love is described as a force that compels the beloved to love in return. The line is often discussed for its beautiful but dangerous fatalism.",
    },
    {
        "phrase": "La gloria di colui che tutto move per l'universo penetra.",
        "source": "Dante, Commedia, Paradiso I, 1-2; Italian original from Kalliope",
        "explanation_zh": "这是《天堂篇》的开篇，把神圣光辉写成穿透宇宙万物的力量；tutto move 与前面“推动太阳和群星”的意象相呼应。",
        "explanation_en": "This opening of Paradiso imagines divine glory as a light that penetrates the whole universe, joining movement, order, and illumination.",
    },
    {
        "phrase": "Considerate la vostra semenza.",
        "source": "Dante, Commedia, Inferno XXVI, 118; Italian original from Kalliope",
        "explanation_zh": "这句可理解为“想想你们的本源”。在尤利西斯的演说中，它引出人应追求德性与知识的名句。",
        "explanation_en": "The line asks listeners to consider their origin or nature. In Ulysses' speech, it prepares the call to pursue virtue and knowledge.",
    },
    {
        "phrase": "Libertà va cercando, ch'è sì cara.",
        "source": "Dante, Commedia, Purgatorio I, 71; Italian original from Kalliope",
        "explanation_zh": "这句说“他正在寻找自由，而自由如此珍贵”。在《炼狱篇》开端，它把自由写成重新上升和净化的核心目标。",
        "explanation_en": "The line says that the traveller is seeking freedom, and that freedom is precious. At the opening of Purgatorio, it frames ascent as a search for liberation.",
    },
    {
        "phrase": "Puro e disposto a salire a le stelle.",
        "source": "Dante, Commedia, Purgatorio XXXIII, 145; Italian original from Kalliope",
        "explanation_zh": "这是《炼狱篇》的结尾，意思是“纯净并准备好上升到群星”。它常被理解为经历修正之后重新拥有前行能力。",
        "explanation_en": "This closing line of Purgatorio means that the speaker is purified and ready to rise to the stars. It suggests restored readiness after difficult correction.",
    },
    {
        "phrase": "Trasumanar significar per verba non si poria.",
        "source": "Dante, Commedia, Paradiso I, 70-71; Italian original from Kalliope",
        "explanation_zh": "这句说“超越人的状态，不能完全用语言说明”。它表达经验超过普通语言边界时的困难。",
        "explanation_en": "The line says that going beyond the human condition cannot fully be expressed in words. It points to experience at the edge of language.",
    },
    {
        "phrase": "Non ragioniam di lor, ma guarda e passa.",
        "source": "Dante, Commedia, Inferno III, 51; Italian original from Kalliope",
        "explanation_zh": "这句大意是“不必谈论他们，只看一眼，然后走过去”。现代语境中常用来提醒自己不要被无意义的噪声拖住。",
        "explanation_en": "The line means, roughly, do not dwell on them; look and pass on. It is often used as advice not to be trapped by empty noise.",
    },
    {
        "phrase": "Ahi serva Italia, di dolore ostello.",
        "source": "Dante, Commedia, Purgatorio VI, 76; Italian original from Kalliope",
        "explanation_zh": "这是但丁对意大利政治分裂的痛切感叹，意思近于“唉，受奴役的意大利，痛苦的居所”。",
        "explanation_en": "This is Dante's bitter lament over political division in Italy, calling the country a dwelling place of sorrow.",
    },
    {
        "phrase": "Tu proverai sì come sa di sale lo pane altrui.",
        "source": "Dante, Commedia, Paradiso XVII, 58-59; Italian original from Kalliope",
        "explanation_zh": "这句说“你会尝到别人的面包有多咸”，常被解读为流亡和寄人篱下的辛酸。",
        "explanation_en": "The line says that one will learn how salty another person's bread tastes, a vivid image of exile and dependence.",
    },
    {
        "phrase": "Per me si va ne la città dolente.",
        "source": "Dante, Commedia, Inferno III, 1; Italian original from Kalliope",
        "explanation_zh": "这是地狱门铭的开头，意为“由我进入痛苦之城”。它比后面的“放弃希望”更像一段严峻旅程的入口。",
        "explanation_en": "This is the opening of the inscription over Hell's gate: through me one enters the city of sorrow. It marks the threshold of a hard journey.",
    },
    {
        "phrase": "Ché la diritta via era smarrita.",
        "source": "Dante, Commedia, Inferno I, 3; Italian original from Kalliope",
        "explanation_zh": "这是开篇迷失场景的核心句，意思是“因为正路已经失去”。它常被用来描述方向感崩塌后的自我觉察。",
        "explanation_en": "This completes the opening image of disorientation: the straight path has been lost. It is a concise image of losing direction.",
    },
    {
        "phrase": "Io non so ben ridir com'i' v'intrai.",
        "source": "Dante, Commedia, Inferno I, 10; Italian original from Kalliope",
        "explanation_zh": "这句说“我也说不清自己怎样走了进去”。它把迷失写成一种渐渐发生、事后才意识到的状态。",
        "explanation_en": "The speaker cannot clearly say how he entered the dark wood. It captures how confusion can happen gradually and only be understood later.",
    },
    {
        "phrase": "A te convien tenere altro viaggio.",
        "source": "Dante, Commedia, Inferno I, 91; Italian original from Kalliope",
        "explanation_zh": "维吉尔告诉但丁必须走“另一条路”。这句话常可理解为：真正的出口不一定是原路返回。",
        "explanation_en": "Virgil tells Dante that he must take another journey. The line suggests that the real way out may not be the way one first expected.",
    },
    {
        "phrase": "Qui si convien lasciare ogne sospetto.",
        "source": "Dante, Commedia, Inferno III, 14; Italian original from Kalliope",
        "explanation_zh": "进入地狱门前，维吉尔要求放下犹疑。这句可用来表达面对艰难任务时必须先停止摇摆。",
        "explanation_en": "Before entering Hell, Virgil tells Dante to leave suspicion behind. It is a line about stopping hesitation before a difficult passage.",
    },
    {
        "phrase": "Ogne viltà convien che qui sia morta.",
        "source": "Dante, Commedia, Inferno III, 15; Italian original from Kalliope",
        "explanation_zh": "这句紧接上一句，意思是“所有怯懦都必须在这里死去”。它强调临界时刻的勇气。",
        "explanation_en": "This follows the call to abandon doubt: every cowardice must die here. It sharpens the threshold into a demand for courage.",
    },
    {
        "phrase": "Vuolsi così colà dove si puote ciò che si vuole.",
        "source": "Dante, Commedia, Inferno III, 95-96; Italian original from Kalliope",
        "explanation_zh": "这是维吉尔多次使用的权威句式，意思是“在能实现意志之处，事情就是如此被愿意的”。它表达不可违抗的天命秩序。",
        "explanation_en": "Virgil uses this formula of authority: so it is willed where what is willed can be done. It signals an order beyond ordinary objection.",
    },
    {
        "phrase": "Fama di loro il mondo esser non lassa.",
        "source": "Dante, Commedia, Inferno III, 49; Italian original from Kalliope",
        "explanation_zh": "这句写那些没有立场者：世界不让他们留下名声。它像一句严厉的警告，提醒人不能总躲在中立里。",
        "explanation_en": "The line describes those who took no stand: the world does not let them have fame. It is a severe warning against empty neutrality.",
    },
    {
        "phrase": "Amor, ch'al cor gentil ratto s'apprende.",
        "source": "Dante, Commedia, Inferno V, 100; Italian original from Kalliope",
        "explanation_zh": "Francesca 叙述爱情时说，爱会迅速抓住高贵的心。这句显示但丁把爱情写得美丽又危险。",
        "explanation_en": "Francesca says that love quickly takes hold of a gentle heart. The line shows love as both beautiful and dangerous.",
    },
    {
        "phrase": "Amor condusse noi ad una morte.",
        "source": "Dante, Commedia, Inferno V, 106; Italian original from Kalliope",
        "explanation_zh": "这句说“爱把我们带向同一场死亡”。它把浪漫叙述转为悲剧后果。",
        "explanation_en": "Love led us to one death. The line turns a romantic story into the stark language of consequence.",
    },
    {
        "phrase": "Nessun maggior dolore che ricordarsi del tempo felice ne la miseria.",
        "source": "Dante, Commedia, Inferno V, 121-123; Italian original from Kalliope",
        "explanation_zh": "这句说，在苦难中回想幸福时光，没有更大的痛苦。它是《地狱篇》中最著名的记忆之痛。",
        "explanation_en": "There is no greater sorrow than remembering happy times in misery. It is one of Inferno's most famous lines about memory and pain.",
    },
    {
        "phrase": "Galeotto fu 'l libro e chi lo scrisse.",
        "source": "Dante, Commedia, Inferno V, 137; Italian original from Kalliope",
        "explanation_zh": "Francesca 说那本书和作者成了撮合者。它常被引用来说明阅读、欲望和行动之间微妙的边界。",
        "explanation_en": "Francesca says the book and its author were the go-between. The line is often cited for the dangerous intimacy between reading and action.",
    },
    {
        "phrase": "E caddi come corpo morto cade.",
        "source": "Dante, Commedia, Inferno V, 142; Italian original from Kalliope",
        "explanation_zh": "这是第五歌结尾，但丁听完 Francesca 的故事后昏倒。句子极短，却把同情的冲击写到身体层面。",
        "explanation_en": "At the end of Canto V, Dante faints after hearing Francesca. The line is brief, physical, and unforgettable.",
    },
    {
        "phrase": "La bocca sollevò dal fiero pasto.",
        "source": "Dante, Commedia, Inferno XXXIII, 1; Italian original from Kalliope",
        "explanation_zh": "这是乌戈利诺故事的开头，画面极强：他从残酷的吞噬中抬起嘴。它代表《地狱篇》晚段的黑暗叙事力量。",
        "explanation_en": "This opens Ugolino's story with a brutal visual image. It is one of Inferno's darkest and most cinematic beginnings.",
    },
    {
        "phrase": "Poscia, più che 'l dolor, poté 'l digiuno.",
        "source": "Dante, Commedia, Inferno XXXIII, 75; Italian original from Kalliope",
        "explanation_zh": "乌戈利诺故事中最可怕的含混句之一，可理解为“后来，饥饿胜过了悲痛”。它的力量来自不说破。",
        "explanation_en": "One of Ugolino's most terrible ambiguous lines: afterward, hunger had more power than grief. Its force lies in what it refuses to spell out.",
    },
    {
        "phrase": "O frati, dissi, che per cento milia perigli siete giunti a l'occidente.",
        "source": "Dante, Commedia, Inferno XXVI, 112-113; Italian original from Kalliope",
        "explanation_zh": "这是尤利西斯演说的开场，呼唤同伴回顾他们穿过的无数危险。它为“追求德性与知识”的名句铺垫。",
        "explanation_en": "This begins Ulysses' speech to his companions, recalling the dangers they have crossed. It prepares the famous call to virtue and knowledge.",
    },
    {
        "phrase": "Tre volte il fé girar con tutte l'acque.",
        "source": "Dante, Commedia, Inferno XXVI, 139; Italian original from Kalliope",
        "explanation_zh": "这是尤利西斯船难场景中的名句，写船被海水三次旋转。它把求知冒险推向悲剧终点。",
        "explanation_en": "In Ulysses' shipwreck, the sea turns the vessel three times. The line carries heroic exploration into tragedy.",
    },
    {
        "phrase": "Dolce color d'oriental zaffiro.",
        "source": "Dante, Commedia, Purgatorio I, 13; Italian original from Kalliope",
        "explanation_zh": "《炼狱篇》开头的天空颜色，像东方蓝宝石一样清甜。它与《地狱篇》的黑暗形成鲜明转场。",
        "explanation_en": "At the opening of Purgatorio, the sky has the sweet color of oriental sapphire. It is a radiant transition out of Inferno's darkness.",
    },
    {
        "phrase": "Che perder tempo a chi più sa più spiace.",
        "source": "Dante, Commedia, Purgatorio III, 78; Italian original from Kalliope",
        "explanation_zh": "这句说，越明智的人越厌恶浪费时间。它是一句非常直接的行动提醒。",
        "explanation_en": "The wiser one is, the more painful it is to waste time. It is one of the poem's cleanest reminders to keep moving.",
    },
    {
        "phrase": "Vien dietro a me, e lascia dir le genti.",
        "source": "Dante, Commedia, Purgatorio V, 13; Italian original from Kalliope",
        "explanation_zh": "维吉尔说：跟着我，让别人去说吧。它常被当作不被议论牵着走的格言。",
        "explanation_en": "Virgil says: follow me, and let people talk. It is often read as counsel against being ruled by gossip.",
    },
    {
        "phrase": "Sta come torre ferma, che non crolla.",
        "source": "Dante, Commedia, Purgatorio V, 14; Italian original from Kalliope",
        "explanation_zh": "这句写人应像坚固的塔一样站稳，不因风声而动摇。它是《炼狱篇》中很适合做每日提醒的句子。",
        "explanation_en": "Stand like a firm tower that does not collapse. It is a compact image of steadiness under pressure.",
    },
    {
        "phrase": "Era già l'ora che volge il disio ai navicanti.",
        "source": "Dante, Commedia, Purgatorio VIII, 1; Italian original from Kalliope",
        "explanation_zh": "这是黄昏时分的名句：那是让航海者转向思乡之情的时刻。它把时间、旅途和乡愁连在一起。",
        "explanation_en": "This famous evening line describes the hour that turns sailors' desire homeward. It joins time, travel, and longing.",
    },
    {
        "phrase": "O vana gloria de l'umane posse!",
        "source": "Dante, Commedia, Purgatorio XI, 91; Italian original from Kalliope",
        "explanation_zh": "这句感叹人间能力和名声的虚浮。它是但丁关于艺术、声誉与时间的一次清醒反思。",
        "explanation_en": "This exclamation calls human glory vain. It is Dante thinking lucidly about art, reputation, and time.",
    },
    {
        "phrase": "Non v'accorgete voi che noi siam vermi nati a formar l'angelica farfalla?",
        "source": "Dante, Commedia, Purgatorio X, 124-125; Italian original from Kalliope",
        "explanation_zh": "这句把人写成将来会形成天使蝴蝶的虫。它用强烈意象表达人的未完成和转化潜能。",
        "explanation_en": "Dante says humans are like worms born to form the angelic butterfly. It is a striking image of incompletion and transformation.",
    },
    {
        "phrase": "Lo maggior don che Dio per sua larghezza fesse creando.",
        "source": "Dante, Commedia, Paradiso V, 19-20; Italian original from Kalliope",
        "explanation_zh": "这句谈自由意志，把它称为上帝创造时赐予的最大礼物。它是《天堂篇》关于自由的核心表达之一。",
        "explanation_en": "This line speaks of free will as the greatest gift God gave in creation. It is one of Paradiso's central statements about freedom.",
    },
    {
        "phrase": "In la sua volontade è nostra pace.",
        "source": "Dante, Commedia, Paradiso III, 85; Italian original from Kalliope",
        "explanation_zh": "这句说“在他的意志中有我们的平安”。它是《天堂篇》中最常被引用的神学句之一。",
        "explanation_en": "In his will is our peace. This is one of Paradiso's most quoted theological lines.",
    },
    {
        "phrase": "Poca favilla gran fiamma seconda.",
        "source": "Dante, Commedia, Paradiso I, 34; Italian original from Kalliope",
        "explanation_zh": "这句说，小火花之后可以跟随大火焰。它适合用来表达微小开端可能引出巨大结果。",
        "explanation_en": "A small spark may be followed by a great flame. It is a compact image of large consequences from small beginnings.",
    },
    {
        "phrase": "O insensata cura de' mortali.",
        "source": "Dante, Commedia, Paradiso XI, 1; Italian original from Kalliope",
        "explanation_zh": "这句感叹凡人的忧虑多么无理。它从天堂视角回看世俗焦虑，带有强烈反讽。",
        "explanation_en": "This line laments the senseless cares of mortals. From Paradiso's height, ordinary anxieties look painfully misplaced.",
    },
    {
        "phrase": "Vergine Madre, figlia del tuo figlio.",
        "source": "Dante, Commedia, Paradiso XXXIII, 1; Italian original from Kalliope",
        "explanation_zh": "这是《天堂篇》终章祈祷的开头，悖论式地称玛利亚为“童贞母亲，你儿子的女儿”。",
        "explanation_en": "This opens the final prayer of Paradiso, calling Mary Virgin Mother and daughter of her son in a powerful theological paradox.",
    },
    {
        "phrase": "Nel suo profondo vidi che s'interna, legato con amore in un volume.",
        "source": "Dante, Commedia, Paradiso XXXIII, 85-86; Italian original from Kalliope",
        "explanation_zh": "但丁在终章异象中看见宇宙万物被爱装订成一卷。它是《神曲》关于统一与爱的最高意象之一。",
        "explanation_en": "In the final vision, Dante sees all things bound by love into one volume. It is one of the poem's supreme images of unity.",
    },
    {
        "phrase": "Ciò che per l'universo si squaderna.",
        "source": "Dante, Commedia, Paradiso XXXIII, 87; Italian original from Kalliope",
        "explanation_zh": "这句承接上一句，指那些在宇宙中分散展开的一切。它与“被爱装订成一卷”的意象构成对照。",
        "explanation_en": "This line refers to what is scattered through the universe. It completes the contrast with all things bound together by love.",
    },
    {
        "phrase": "O luce etterna che sola in te sidi.",
        "source": "Dante, Commedia, Paradiso XXXIII, 124; Italian original from Kalliope",
        "explanation_zh": "终章中但丁呼唤“永恒之光”。它把《天堂篇》的视觉和神学主题推到顶点。",
        "explanation_en": "Near the end, Dante addresses the eternal light. It brings Paradiso's visual and theological language to its height.",
    },
    {
        "phrase": "A l'alta fantasia qui mancò possa.",
        "source": "Dante, Commedia, Paradiso XXXIII, 142; Italian original from Kalliope",
        "explanation_zh": "这句说，在这里崇高的想象力也失去了力量。它承认语言和想象在终极经验前的限度。",
        "explanation_en": "Here the high imagination failed in power. The line acknowledges the limit of imagination before ultimate vision.",
    },
]
MARINE_CONTEXT_TERMS = [
    "ocean",
    "marine",
    "sea",
    "seawater",
    "coastal",
    "coast",
    "estuary",
    "shelf",
    "basin",
    "atlantic",
    "pacific",
    "indian ocean",
    "southern ocean",
    "arabian sea",
    "mediterranean",
    "gulf",
    "strait",
    "bgc-argo",
    "biogeochemical argo",
    "argo float",
    "ocean colour",
    "ocean color",
    "ocean optics",
    "marine heatwave",
    "marine heatwaves",
]
DOMAIN_CONTEXT_TERMS = [
    "bgc-argo",
    "biogeochemical argo",
    "carbon pump",
    "carbon export",
    "net community production",
    "particulate organic carbon",
    "poc",
    "microbial carbon",
    "bacterioplankton",
    "microbiome",
    "phytoplankton",
    "chlorophyll",
    "deep chlorophyll maximum",
    "subsurface chlorophyll",
    "primary production",
    "net primary production",
    "ocean colour",
    "ocean color",
    "ocean optics",
    "aquatic optics",
    "bio-optic",
    "bio-optical",
    "inherent optical properties",
    "apparent optical properties",
    "apparent optical property",
    "iop",
    "iops",
    "aop",
    "aops",
    "in situ",
    "in-situ",
    "optical absorption",
    "absorption coefficient",
    "absorption coefficients",
    "phytoplankton absorption",
    "particulate absorption",
    "particulate absorption coefficient",
    "particulate absorption coefficients",
    "particle absorption",
    "cdom absorption",
    "colored dissolved organic matter",
    "coloured dissolved organic matter",
    "backscatter",
    "backscattering",
    "backscattering coefficient",
    "particulate backscattering",
    "particle backscattering",
    "angular scattering",
    "volume scattering",
    "very small particles",
    "bbp",
    "vsp",
    "remote sensing reflectance",
    "rrs",
    "diffuse attenuation",
    "water-leaving radiance",
    "ocean lidar",
    "remote sensing",
    "hyperspectral",
    "marine heatwave",
    "marine heatwaves",
]


@dataclass
class Paper:
    title: str
    authors: str
    journal: str
    published: str
    published_month: str
    published_month_zh: str
    doi: str
    url: str
    abstract: str
    group_zh: str
    group_en: str
    tags: str
    summary_zh: str
    summary_en: str
    score: int
    rank: int


def http_json(url: str, retries: int = 2) -> dict:
    for attempt in range(retries):
        request = urllib.request.Request(url, headers={"User-Agent": USER_AGENT})
        try:
            with urllib.request.urlopen(request, timeout=15) as response:
                return json.loads(response.read().decode("utf-8"))
        except (urllib.error.URLError, TimeoutError) as exc:
            if attempt == retries - 1:
                raise exc
            time.sleep(1.5 * (attempt + 1))
    raise RuntimeError("unreachable")


def text_value(value) -> str:
    if isinstance(value, list):
        return str(value[0]) if value else ""
    return str(value or "")


def clean_text(value: str) -> str:
    value = re.sub(r"<[^>]+>", " ", html.unescape(value or ""))
    value = re.sub(r"\s+", " ", value).strip()
    return value


def month_from_parts(parts: list[int], fallback: str = "") -> tuple[str, str, str]:
    if len(parts) >= 2:
        month = f"{parts[0]:04d}-{parts[1]:02d}"
        published = "-".join(f"{x:02d}" if i else f"{x:04d}" for i, x in enumerate(parts[:3]))
        return published, month, month
    if len(parts) == 1:
        year = str(parts[0])
        return year, f"{year} (month not verified)", f"{year}（月份未核准）"
    return fallback, fallback, fallback


def published_parts(item: dict) -> list[int]:
    for key in ("published-print", "published-online", "published", "created"):
        parts = item.get(key, {}).get("date-parts")
        if parts and parts[0]:
            return parts[0]
    return []


def is_future_publication(parts: list[int], today: date) -> bool:
    if len(parts) >= 3:
        return date(parts[0], parts[1], parts[2]) > today
    if len(parts) >= 2:
        return (parts[0], parts[1]) > (today.year, today.month)
    if len(parts) == 1:
        return parts[0] > today.year
    return False


def authors_from_item(item: dict) -> str:
    names = []
    for author in item.get("author", [])[:8]:
        given = author.get("given", "").strip()
        family = author.get("family", "").strip()
        name = " ".join(part for part in [given, family] if part)
        if name:
            names.append(name)
    if len(item.get("author", [])) > 8:
        names.append("et al.")
    return "; ".join(names) or "Authors not available"


def author_names_from_item(item: dict) -> list[str]:
    names = []
    for author in item.get("author", []):
        given = author.get("given", "").strip()
        family = author.get("family", "").strip()
        name = " ".join(part for part in [given, family] if part)
        if name:
            names.append(name)
    return names


def normalize_person_name(value: str) -> str:
    value = unicodedata.normalize("NFKD", value or "")
    value = "".join(ch for ch in value if not unicodedata.combining(ch))
    return re.sub(r"[^a-z0-9]+", " ", value.lower()).strip()


def person_signature(value: str, *, reverse_two_part: bool = False) -> str:
    tokens = normalize_person_name(value).split()
    if reverse_two_part and len(tokens) == 2:
        tokens = list(reversed(tokens))
    if not tokens:
        return ""
    family = tokens[-1]
    initials = "".join(token[0] for token in tokens[:-1] if token)
    return f"{family}:{initials}"


def person_name_keys(value: str) -> set[str]:
    normalized = normalize_person_name(value)
    if not normalized:
        return set()
    keys = {normalized.replace(" ", ""), person_signature(value)}
    if len(normalized.split()) == 2:
        keys.add(person_signature(value, reverse_two_part=True))
    return {key for key in keys if key}


@lru_cache(maxsize=1)
def focused_team_config() -> dict:
    if not FOCUSED_TEAM_AUTHORS_PATH.exists():
        return {"authors": [], "topic_keywords": []}
    try:
        return json.loads(FOCUSED_TEAM_AUTHORS_PATH.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        print(f"Focused team author file could not be read: {exc}", file=sys.stderr)
        return {"authors": [], "topic_keywords": []}


def focused_team_author_records() -> list[dict]:
    return focused_team_config().get("authors", [])


def focused_team_query_text() -> str:
    keywords = focused_team_config().get("topic_keywords", [])
    return " ".join(str(keyword) for keyword in keywords[:10])


@lru_cache(maxsize=1)
def focused_team_author_keys() -> set[str]:
    keys: set[str] = set()
    for record in focused_team_author_records():
        keys.update(person_name_keys(record.get("name", "")))
    return keys


def is_focused_team_item(item: dict) -> bool:
    team_keys = focused_team_author_keys()
    if not team_keys:
        return False
    for name in author_names_from_item(item):
        if person_name_keys(name) & team_keys:
            return True
    return False


def best_journal(item: dict) -> str:
    titles = item.get("container-title") or []
    return text_value(titles).strip()


def normalize_journal(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", value.lower()).strip()


def canonical_journal(journal: str) -> str:
    low = journal.lower()
    normalized = normalize_journal(journal)
    for candidate in PRIORITY_JOURNALS + NATURE_SCIENCE_HINTS:
        candidate_low = candidate.lower()
        candidate_normalized = normalize_journal(candidate)
        if candidate_low in {"nature", "science"}:
            if low == candidate_low or low.startswith(f"{candidate_low} "):
                return candidate
            continue
        if candidate_normalized in normalized:
            return candidate
    return journal


def is_priority_journal(journal: str) -> bool:
    canonical = canonical_journal(journal)
    return canonical in JOURNAL_RANK or canonical in NATURE_SCIENCE_HINTS


def group_for(journal: str, focused_team: bool = False) -> tuple[str, str, int]:
    low = journal.lower()
    if journal in NATURE_SCIENCE_HINTS and "nature" in low:
        return "Nature 系列", "Nature series", 0
    if journal in NATURE_SCIENCE_HINTS and (low == "science" or low.startswith("science ")):
        return "Science 系列", "Science series", 1
    if journal in JOURNAL_RANK:
        return "重点期刊：按影响力和相关性排序", "Key journals: ordered by impact and relevance", 2
    if focused_team:
        return "重点关注团队", "Focused team", 3
    return "其他相关期刊：按主题相关性补充", "Other relevant journals: topical supplements", 4


def relevance_score(title: str, abstract: str) -> int:
    haystack = f"{title} {abstract}".lower()
    return sum(weight for term, weight in RELEVANCE_TERMS.items() if term_present(term, haystack))


def term_present(term: str, haystack: str) -> bool:
    if term in {"iop", "iops", "aop", "aops", "bbp", "aph", "adg", "vsp", "rrs", "kd"}:
        return bool(re.search(rf"\b{re.escape(term)}\b", haystack))
    return term in haystack


def has_marine_context(title: str, abstract: str, journal: str) -> bool:
    haystack = f"{title} {abstract} {journal}".lower()
    return any(re.search(rf"\b{re.escape(term)}\b", haystack) for term in MARINE_CONTEXT_TERMS)


def has_domain_context(title: str, abstract: str, journal: str = "") -> bool:
    haystack = f"{title} {abstract} {journal}".lower()
    return any(re.search(rf"\b{re.escape(term)}\b", haystack) for term in DOMAIN_CONTEXT_TERMS)


def tags_for(title: str, abstract: str) -> str:
    haystack = f"{title} {abstract}".lower()
    tags = []
    checks = [
        ("BGC-Argo", ["bgc-argo", "biogeochemical argo", "argo float"]),
        ("carbon pump", ["carbon pump", "carbon export", "net community production", "poc"]),
        ("phytoplankton", ["phytoplankton", "chlorophyll", "surface ocean chlorophyll"]),
        ("marine heatwaves", ["marine heatwave", "marine heatwaves"]),
        ("ocean colour", ["ocean colour", "ocean color", "remote sensing"]),
        ("ocean optics", ["ocean optics", "aquatic optics", "inherent optical", "apparent optical", "remote sensing reflectance", "rrs", "water-leaving radiance", "aop", "aops"]),
        ("absorption", ["optical absorption", "absorption coefficient", "absorption coefficients", "phytoplankton absorption", "particulate absorption", "particulate absorption coefficient", "particulate absorption coefficients", "particle absorption", "cdom absorption", "colored dissolved organic matter", "coloured dissolved organic matter", "aph", "adg"]),
        ("backscattering", ["backscatter", "backscattering", "backscattering coefficient", "particulate backscattering", "particle backscattering", "angular scattering", "volume scattering", "bbp", "vsp"]),
        ("bio-optics", ["bio-optic", "bio-optical", "iop", "iops", "aop", "aops", "optical", "diffuse attenuation", "ocean lidar"]),
        ("microbial carbon", ["microbial carbon", "microbial", "dissolved organic"]),
        ("vertical structure", ["vertical", "subsurface", "deep chlorophyll maximum", "dcm"]),
    ]
    for label, terms in checks:
        if any(term_present(term, haystack) for term in terms):
            tags.append(label)
    return "; ".join(tags[:5]) or "ocean biogeochemistry"


def sentence_split(text: str) -> list[str]:
    parts = re.split(r"(?<=[.!?])\s+", text)
    return [part.strip() for part in parts if part.strip()]


def topic_label(title: str, tags: str) -> tuple[str, str]:
    haystack = f"{title} {tags}".lower()
    if "bgc-argo" in haystack and "oxygen-deficient" in haystack:
        return "氧亏区中的氮-碳循环变化", "nitrogen-carbon cycling in oxygen-deficient waters"
    if "bgc-argo" in haystack:
        return "BGC-Argo 剖面约束的生物地球化学变化", "biogeochemical variability constrained by BGC-Argo profiles"
    if "marine heatwave" in haystack or "marine heatwaves" in haystack:
        return "海洋热浪对生态结构和碳循环的影响", "the ecological and carbon-cycle impacts of marine heatwaves"
    if any(term in haystack for term in ["ocean optics", "aquatic optics", "bio-optic", "bio-optical", "backscatter", "backscattering", "scattering", "absorption", "remote sensing reflectance", "rrs", "iop", "iops", "aop", "aops", "cdom", "diffuse attenuation", "ocean lidar"]):
        return "海洋生物光学和海洋光学参数、方法与应用", "marine bio-optical and ocean-optics properties, methods, and applications"
    if "surface ocean chlorophyll" in haystack:
        return "表层海洋叶绿素长期变化及其气候驱动", "long-term changes in surface-ocean chlorophyll and their climatic drivers"
    if "ocean colour" in haystack or "ocean color" in haystack or "remote sensing" in haystack:
        return "海色遥感中的浮游植物和光学信号变化", "phytoplankton and optical signals in ocean-colour remote sensing"
    if "carbon export" in haystack or "carbon pump" in haystack or "organic carbon" in haystack:
        return "有机碳输出和生物碳泵过程", "organic carbon export and biological carbon-pump processes"
    if "phytoplankton" in haystack:
        return "浮游植物群落、生理状态和生产力变化", "changes in phytoplankton communities, physiology, and productivity"
    if "bio-optic" in haystack or "backscatter" in haystack or "absorption" in haystack or "lidar" in haystack:
        return "海洋生物光学参数及其观测方法", "marine bio-optical properties and observing methods"
    return "海洋生物地球化学过程", "ocean biogeochemical processes"


def infer_data_source(title: str, tags: str, journal: str) -> tuple[str, str]:
    haystack = f"{title} {tags} {journal}".lower()
    if "bgc-argo" in haystack or "argo" in haystack:
        return (
            "数据基础是 BGC-Argo 浮标剖面及相关生物地球化学观测，可用于连接垂向结构和过程变化。",
            "The data basis is BGC-Argo float profiles and related biogeochemical observations, which connect vertical structure with process variability.",
        )
    if "pace" in haystack:
        return (
            "数据或应用场景围绕 NASA PACE/OCI 高光谱海色观测、产品验证和社区使用展开。",
            "The data or application setting centres on NASA PACE/OCI hyperspectral ocean-colour observations, product validation, and community use.",
        )
    if "viirs" in haystack:
        return (
            "数据上使用 VIIRS 顶层大气卫星观测，面向近岸复杂水体的叶绿素 a 估计。",
            "The data source is top-of-atmosphere VIIRS satellite observation, used for chlorophyll-a estimation in optically complex coastal waters.",
        )
    if any(term in haystack for term in ["satellite", "ocean color", "ocean colour", "chlorophyll", "remote sensing"]):
        return (
            "数据上主要面向卫星海色、叶绿素或遥感反射率记录，并在需要时结合原位或剖面观测进行约束。",
            "The data emphasis is satellite ocean-colour, chlorophyll, or remote-sensing reflectance records, with in-situ or profile constraints where relevant.",
        )
    if any(term in haystack for term in ["sediment trap", "optics", "backscattering", "backscatter", "bio-optic", "bio-optical", "absorption", "inherent optical", "apparent optical", "remote sensing reflectance", "rrs", "diffuse attenuation", "water-leaving radiance", "lidar"]):
        return (
            "数据或方法上围绕沉积物捕获器、原位光学、生物光学参数或光谱观测展开。",
            "The data or method emphasis is sediment traps, in-situ optics, bio-optical parameters, ocean-optics products, or spectral optical observation.",
        )
    if any(term in haystack for term in ["review", "annual review", "modeling", "campaign", "hackweek"]):
        return (
            "这类论文主要综合已有观测、理论、模型、外场设计或社区训练材料，而不是发布一个新的单一观测数据集。",
            "This type of paper synthesizes existing observations, theory, models, field-design material, or community-training workflows rather than publishing one new observational dataset.",
        )
    return (
        "研究结合观测、模型或实验结果来刻画海洋生态和生物地球化学过程。",
        "The study combines observations, models, or experimental results to characterize ocean ecological and biogeochemical processes.",
    )


def infer_action(title: str, tags: str) -> tuple[str, str]:
    haystack = f"{title} {tags}".lower()
    if any(term in haystack for term in ["machine learning", "deep learning", "neural", "emulator"]):
        return (
            "方法上使用机器学习、深度学习或神经模拟器来提取信号、反演变量或识别驱动机制。",
            "Methodologically, it uses machine learning, deep learning, or a neural emulator to extract signals, retrieve variables, or identify drivers.",
        )
    if any(term in haystack for term in ["retrieval", "estimation", "estimating"]):
        return (
            "研究重点是建立或评估反演/估计框架，把观测信号转化为浮游植物、叶绿素、吸收或碳通量等变量。",
            "The main task is to build or assess a retrieval or estimation framework that converts observations into variables such as phytoplankton, chlorophyll, absorption, backscattering, IOPs, or carbon flux.",
        )
    if any(term in haystack for term in ["dataset", "data record", "data with", "surveys"]):
        return (
            "研究重点是整理、质控或构建可复用数据产品，使后续跨区域和长期分析更可靠。",
            "The main task is to assemble, quality-control, or build a reusable data product so that later cross-regional and long-term analyses are more robust.",
        )
    if any(term in haystack for term in ["review", "modeling", "lifting the lid"]):
        return (
            "研究主要通过综述、模型框架或概念整合来梳理关键过程和不确定性来源。",
            "The paper mainly uses review, modelling, or conceptual synthesis to organize key processes and sources of uncertainty.",
        )
    return (
        "研究主要分析变化趋势、事件响应、驱动机制或方法表现，并把观测结果与生态和生物地球化学过程联系起来。",
        "The main task is to analyse trends, event responses, drivers, or method performance, linking observations with ecological or biogeochemical processes.",
    )


def infer_result(title: str, tags: str) -> tuple[str, str]:
    haystack = f"{title} {tags}".lower()
    if any(term in haystack for term in ["declin", "increase", "threefold", "1 °c", "unprecedented"]):
        return (
            "核心结果是识别出上升、下降或异常事件信号，并把这些变化放进气候变暖、生产力或碳循环背景下解释。",
            "The core result is an identified increase, decline, or extreme-event signal interpreted in the context of warming, productivity, or carbon cycling.",
        )
    if any(term in haystack for term in ["accuracy", "comparison", "performance"]):
        return (
            "核心结果是给出产品或方法表现的评估，为判断适用水体、误差来源和后续改进提供依据。",
            "The main outcome is an assessment of product or method performance, helping identify suitable waters, error sources, and paths for improvement.",
        )
    if any(term in haystack for term in ["framework", "approach", "design", "dataset", "data record"]):
        return (
            "核心实现是提出可复用的框架、流程或数据产品，使类似观测、反演、验证或趋势分析可以更系统地开展。",
            "The main contribution is a reusable framework, workflow, or data product that makes similar observation, retrieval, validation, or trend-analysis tasks more systematic.",
        )
    if "heatwave" in haystack or "heatwaves" in haystack:
        return (
            "核心结果帮助说明海洋热浪如何通过强度、持续时间或垂向结构影响生态和碳循环过程。",
            "The core result helps explain how marine heatwaves affect ecological and carbon-cycle processes through intensity, duration, or vertical structure.",
        )
    return (
        "核心结果或实现为该主题提供了新的观测约束、方法基准或过程解释。",
        "The main result or implementation provides new observational constraints, methodological benchmarks, or process interpretation for this topic.",
    )


def summary_en(title: str, abstract: str, tags: str) -> str:
    topic_zh, topic_en = topic_label(title, tags)
    data_zh, data_en = infer_data_source(title, tags, "")
    action_zh, action_en = infer_action(title, tags)
    result_zh, result_en = infer_result(title, tags)
    return " ".join(
        [
            f"The study examines {topic_en}.",
            data_en,
            action_en,
            result_en,
            "It is useful for interpreting observations, comparing methods, and designing follow-up analyses in marine biogeochemistry.",
        ]
    )


def summary_zh(title: str, abstract: str, tags: str) -> str:
    topic_zh, topic_en = topic_label(title, tags)
    data_zh, data_en = infer_data_source(title, tags, "")
    action_zh, action_en = infer_action(title, tags)
    result_zh, result_en = infer_result(title, tags)
    return " ".join(
        [
            f"这项研究关注{topic_zh}。",
            data_zh,
            action_zh,
            result_zh,
            "它有助于解释观测现象、比较方法差异，并为后续海洋生物地球化学分析提供线索。",
        ]
    )


def crossref_query(query: str, from_date: date, rows: int = 20) -> list[dict]:
    params = {
        "query.bibliographic": query,
        "filter": f"from-created-date:{from_date.isoformat()},type:journal-article",
        "sort": "created",
        "order": "desc",
        "rows": str(rows),
        "select": "DOI,title,author,container-title,published,published-print,published-online,created,URL,abstract,type",
    }
    url = "https://api.crossref.org/works?" + urllib.parse.urlencode(params)
    data = http_json(url)
    return data.get("message", {}).get("items", [])


def crossref_author_query(author_name: str, from_date: date, rows: int = 12) -> list[dict]:
    params = {
        "query.author": author_name,
        "query.bibliographic": focused_team_query_text(),
        "filter": f"from-created-date:{from_date.isoformat()},type:journal-article",
        "sort": "created",
        "order": "desc",
        "rows": str(rows),
        "select": "DOI,title,author,container-title,published,published-print,published-online,created,URL,abstract,type",
    }
    url = "https://api.crossref.org/works?" + urllib.parse.urlencode(params)
    data = http_json(url)
    return data.get("message", {}).get("items", [])


def published_sort_value(value: str) -> int:
    parts = [int(part) for part in re.findall(r"\d+", value or "")[:3]]
    if not parts:
        return 0
    year = parts[0]
    month = parts[1] if len(parts) > 1 else 1
    day = parts[2] if len(parts) > 2 else 1
    return year * 10000 + month * 100 + day


def existing_dois() -> set[str]:
    dois = set()
    if DATA_PATH.exists():
        text = DATA_PATH.read_text(encoding="utf-8")
        dois.update(match.lower() for match in re.findall(r'doi:\s*"([^"]+)"', text))
    if HISTORY_PATH.exists():
        with HISTORY_PATH.open("r", encoding="utf-8-sig", newline="") as handle:
            for row in csv.DictReader(handle):
                doi = (row.get("doi") or "").strip().lower()
                if doi:
                    dois.add(doi)
    return dois


def paper_from_crossref_item(item: dict, today: date, require_focused_team: bool = False) -> Paper | None:
    doi = (item.get("DOI") or "").strip()
    if not doi:
        return None
    title = clean_text(text_value(item.get("title")))
    journal = best_journal(item)
    canonical = canonical_journal(journal)
    abstract = clean_text(item.get("abstract", ""))
    if title.lower().startswith(EXCLUDED_TITLE_PREFIXES):
        return None
    if not has_marine_context(title, abstract, journal):
        return None
    if not has_domain_context(title, abstract, journal):
        return None
    parts = published_parts(item)
    if is_future_publication(parts, today):
        return None
    focused_team = is_focused_team_item(item)
    if require_focused_team and not focused_team:
        return None
    score = relevance_score(title, abstract)
    min_score = 5 if is_priority_journal(canonical) else (4 if focused_team else 7)
    if score < min_score:
        return None
    group_zh, group_en, group_rank = group_for(canonical, focused_team)
    published, month, month_zh = month_from_parts(parts)
    tags = tags_for(title, abstract)
    within_group_rank = JOURNAL_RANK.get(canonical, max(0, 500 - score))
    rank = group_rank * 1000 + within_group_rank
    return Paper(
        title=title,
        authors=authors_from_item(item),
        journal=canonical,
        published=published,
        published_month=month,
        published_month_zh=month_zh,
        doi=doi,
        url=item.get("URL") or f"https://doi.org/{doi}",
        abstract=abstract,
        group_zh=group_zh,
        group_en=group_en,
        tags=tags,
        summary_zh=summary_zh(title, abstract, tags),
        summary_en=summary_en(title, abstract, tags),
        score=score,
        rank=rank,
    )


def collect_candidates(lookback_days: int, max_papers: int) -> list[Paper]:
    today = datetime.now(TZ).date()
    from_date = datetime.now(TZ).date() - timedelta(days=lookback_days)
    seen = existing_dois()
    candidates: dict[str, Paper] = {}

    for query in TOPIC_TERMS:
        if len(candidates) >= max_papers * 8:
            break
        try:
            items = crossref_query(query, from_date, rows=100)
        except Exception as exc:
            print(f"Crossref query failed: {query}: {exc}", file=sys.stderr)
            continue
        for item in items:
            doi = (item.get("DOI") or "").strip()
            if not doi or doi.lower() in seen or doi.lower() in candidates:
                continue
            paper = paper_from_crossref_item(item, today)
            if paper:
                candidates[doi.lower()] = paper
        time.sleep(0.05)

    author_rows = max(1, int(os.getenv("FOCUSED_TEAM_AUTHOR_ROWS", "12")))
    author_limit = max(0, int(os.getenv("FOCUSED_TEAM_AUTHOR_LIMIT", "0")))
    author_records = focused_team_author_records()
    if author_limit:
        author_records = author_records[:author_limit]
    author_workers = max(1, int(os.getenv("FOCUSED_TEAM_AUTHOR_WORKERS", "6")))

    def fetch_author_items(record: dict) -> tuple[str, list[dict], Exception | None]:
        author_name = (record.get("name") or "").strip()
        if not author_name:
            return "", [], None
        try:
            return author_name, crossref_author_query(author_name, from_date, rows=author_rows), None
        except Exception as exc:
            return author_name, [], exc

    if author_records:
        with concurrent.futures.ThreadPoolExecutor(max_workers=author_workers) as executor:
            futures = [executor.submit(fetch_author_items, record) for record in author_records]
            for index, future in enumerate(concurrent.futures.as_completed(futures), 1):
                author_name, items, exc = future.result()
                if exc:
                    print(f"Crossref focused-team author query failed: {author_name}: {exc}", file=sys.stderr)
                    continue
                for item in items:
                    doi = (item.get("DOI") or "").strip()
                    if not doi or doi.lower() in seen or doi.lower() in candidates:
                        continue
                    paper = paper_from_crossref_item(item, today, require_focused_team=True)
                    if paper:
                        candidates[doi.lower()] = paper
                if index % 50 == 0:
                    print(f"Focused-team author queries checked: {index}/{len(author_records)}")

    return sorted(
        candidates.values(),
        key=lambda paper: (paper.rank, -paper.score, -published_sort_value(paper.published)),
    )[:max_papers]


def q(value: str) -> str:
    return json.dumps(value or "", ensure_ascii=False)


def previous_italian_uses(issue_date: str) -> dict[str, str]:
    if not DATA_PATH.exists():
        return {}
    text = DATA_PATH.read_text(encoding="utf-8-sig")
    pattern = re.compile(
        r'(?ms)^- date: "([^"]+)"\n.*?^  italian_phrase: "((?:\\.|[^"])*)"',
    )
    uses: dict[str, str] = {}
    for date_text, phrase_json in pattern.findall(text):
        if date_text >= issue_date:
            continue
        try:
            phrase = json.loads(f'"{phrase_json}"')
        except json.JSONDecodeError:
            phrase = phrase_json
        if date_text > uses.get(phrase, ""):
            uses[phrase] = date_text
    return uses


def dante_card_for(issue_date: str) -> dict[str, str]:
    try:
        start = datetime.strptime(issue_date, "%Y-%m-%d").date().toordinal() % len(DANTE_CARDS)
    except ValueError:
        start = 0
    used = previous_italian_uses(issue_date)
    ordered_cards = DANTE_CARDS[start:] + DANTE_CARDS[:start]
    for card in ordered_cards:
        if card["phrase"] not in used:
            return card
    return min(
        ordered_cards,
        key=lambda card: (used.get(card["phrase"], ""), card["phrase"]),
    )


def issue_block(today: str, papers: list[Paper]) -> str:
    generated_at = datetime.now(TZ).strftime("%Y-%m-%d %H:%M %Z")
    title_zh = "每日论文推送：BGC-Argo、海色/海洋光学、海洋热浪与碳泵"
    title_en = "Daily Paper Push: BGC-Argo, ocean colour/ocean optics, marine heatwaves and carbon pump"
    summary_zh = f"本期由 GitHub Actions 自动检索生成：Nature/Science 系列优先，其次是用户指定重点期刊，再补充重点关注团队的新论文，最后纳入其他相关期刊；历史去重后保留 {len(papers)} 篇，不超过每日 50 篇上限。"
    summary_en = f"This issue was generated automatically by GitHub Actions: Nature and Science series first, then the user-defined priority journals, then new papers from the focused team, followed by other relevant journals as topical supplements. After deduplication, {len(papers)} papers remain, below the daily limit of 50."
    trend_zh = "本期重点关注 BGC-Argo、海色遥感/海洋光学、海洋热浪、浮游植物垂向结构和碳泵过程。筛选逻辑不再只限于重点期刊；当高影响力期刊当天新增较少时，会额外检索重点关注团队作者的新论文，并用海洋、海色/光学和碳循环关键词过滤，再从其他相关期刊补充候选论文。"
    trend_en = "This issue focuses on BGC-Argo, ocean-colour remote sensing, ocean optics, marine heatwaves, vertical phytoplankton structure and carbon-pump processes. The selection is no longer limited to priority journals; when few high-impact papers are newly available, the workflow also checks focused-team authors and filters those papers with ocean, ocean-colour/optics, and carbon-cycle keywords before adding other relevant journals as supplements."
    docx_name = f"daily_paper_push_{today}.docx"
    dante = dante_card_for(today)

    lines = [
        f"- date: {q(today)}",
        f"  generated_at: {q(generated_at)}",
        f"  title: {q(title_zh)}",
        f"  title_zh: {q(title_zh)}",
        f"  title_en: {q(title_en)}",
        f"  summary: {q(summary_zh)}",
        f"  summary_zh: {q(summary_zh)}",
        f"  summary_en: {q(summary_en)}",
        f"  trend_zh: {q(trend_zh)}",
        f"  trend_en: {q(trend_en)}",
        f"  docx: {q('/files/paper-push/' + docx_name)}",
        "  figure:",
        f"  italian_phrase: {q(dante['phrase'])}",
        f"  italian_source: {q(dante['source'])}",
        f"  italian_explanation_zh: {q(dante['explanation_zh'])}",
        f"  italian_explanation_en: {q(dante['explanation_en'])}",
        "  papers:",
    ]
    for paper in papers:
        lines.append(f"    - title: {q(paper.title)}")
        for key, value in [
            ("authors", paper.authors),
            ("journal", paper.journal),
            ("published", paper.published),
            ("published_month", paper.published_month),
            ("published_month_zh", paper.published_month_zh),
            ("doi", paper.doi),
            ("url", paper.url),
            ("group", paper.group_zh),
            ("group_zh", paper.group_zh),
            ("group_en", paper.group_en),
            ("tags", paper.tags),
            ("tags_zh", paper.tags),
            ("tags_en", paper.tags),
            ("summary", paper.summary_zh),
            ("summary_zh", paper.summary_zh),
            ("summary_en", paper.summary_en),
        ]:
            lines.append(f"      {key}: {q(value)}")
    return "\n".join(lines) + "\n"


def issue_exists(issue_date: str) -> bool:
    if not DATA_PATH.exists():
        return False
    text = DATA_PATH.read_text(encoding="utf-8-sig")
    return bool(re.search(rf'(?m)^- date: "{re.escape(issue_date)}"$', text))


def latest_previous_issue_date(issue_date: str) -> str:
    if not DATA_PATH.exists():
        return ""
    text = DATA_PATH.read_text(encoding="utf-8-sig")
    dates = re.findall(r'(?m)^- date: "([^"]+)"$', text)
    previous = [date_text for date_text in dates if date_text < issue_date]
    return previous[0] if previous else (dates[0] if dates else "")


def no_update_issue_block(today: str, previous_date: str) -> str:
    generated_at = datetime.now(TZ).strftime("%Y-%m-%d %H:%M %Z")
    title_zh = "每日论文推送：今日暂无新增"
    title_en = "Daily Paper Push: no new papers captured today"
    if previous_date:
        summary_zh = f"今日尚未捕捉到符合更新规则的新论文，论文列表与 {previous_date} 保持一致。"
        summary_en = f"No new papers matching the update rules were captured today; the paper list remains unchanged from {previous_date}."
    else:
        summary_zh = "今日尚未捕捉到符合更新规则的新论文。"
        summary_en = "No new papers matching the update rules were captured today."
    dante = dante_card_for(today)

    return "\n".join(
        [
            f"- date: {q(today)}",
            f"  generated_at: {q(generated_at)}",
            "  no_update: true",
            f"  previous_date: {q(previous_date)}",
            f"  title: {q(title_zh)}",
            f"  title_zh: {q(title_zh)}",
            f"  title_en: {q(title_en)}",
            f"  summary: {q(summary_zh)}",
            f"  summary_zh: {q(summary_zh)}",
            f"  summary_en: {q(summary_en)}",
            "  docx: \"\"",
            "  figure:",
            f"  italian_phrase: {q(dante['phrase'])}",
            f"  italian_source: {q(dante['source'])}",
            f"  italian_explanation_zh: {q(dante['explanation_zh'])}",
            f"  italian_explanation_en: {q(dante['explanation_en'])}",
            "  papers: []",
        ]
    ) + "\n"


def replace_or_prepend_issue(today: str, block: str) -> None:
    existing = DATA_PATH.read_text(encoding="utf-8") if DATA_PATH.exists() else ""
    pattern = re.compile(rf'(?ms)^- date: "{re.escape(today)}"\n.*?(?=^- date: "|\Z)')
    if pattern.search(existing):
        updated = pattern.sub(block, existing).rstrip() + "\n"
    else:
        updated = block + ("\n" + existing if existing.strip() else "")
    DATA_PATH.write_text(updated, encoding="utf-8")


def set_font(run, name: str = "Microsoft YaHei") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def write_docx(today: str, papers: list[Paper]) -> None:
    out_dir = ROOT / "files" / "paper-push"
    out_dir.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(31, 77, 120)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 11.5, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("每日论文推送")
    set_font(r)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(23, 50, 77)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"检索日期：{today} | GitHub Actions 自动生成")
    set_font(r)
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(85, 105, 125)

    doc.add_heading("今日总览", level=1)
    p = doc.add_paragraph()
    r = p.add_run(f"历史去重后今日新增 {len(papers)} 篇。排序规则：Nature 系列、Science 系列、其余重点期刊、重点关注团队、其他相关补充论文。")
    set_font(r)

    current_group = None
    doc.add_heading("论文速读", level=1)
    for index, paper in enumerate(papers, 1):
        if paper.group_zh != current_group:
            current_group = paper.group_zh
            doc.add_heading(current_group, level=2)
        p = doc.add_paragraph(style="Heading 3")
        r = p.add_run(f"{index}. {paper.title}")
        set_font(r)
        r.bold = True
        for label, value in [
            ("作者", paper.authors),
            ("期刊", paper.journal),
            ("发表月份", paper.published_month_zh),
            ("DOI", paper.doi),
        ]:
            p = doc.add_paragraph()
            r = p.add_run(f"{label}：{value}")
            set_font(r)
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(85, 105, 125)
        p = doc.add_paragraph()
        r = p.add_run(f"关键词：{paper.tags}")
        set_font(r)
        r.bold = True
        r.font.size = Pt(9.5)
        p = doc.add_paragraph()
        r = p.add_run(paper.summary_zh)
        set_font(r)
        p = doc.add_paragraph()
        r = p.add_run(f"链接：{paper.url}")
        set_font(r)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(5, 99, 193)

    doc.save(out_dir / f"daily_paper_push_{today}.docx")


def update_history(today: str, papers: list[Paper]) -> None:
    rows = []
    if HISTORY_PATH.exists():
        with HISTORY_PATH.open("r", encoding="utf-8-sig", newline="") as handle:
            rows = list(csv.DictReader(handle))
    seen = {(row.get("doi") or "").lower() for row in rows}
    for paper in papers:
        if paper.doi.lower() in seen:
            continue
        rows.append({
            "date": today,
            "title": paper.title,
            "authors": paper.authors,
            "published_month": paper.published_month_zh,
            "doi": paper.doi,
            "url": paper.url,
            "journal": paper.journal,
        })
    with HISTORY_PATH.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["date", "title", "authors", "published_month", "doi", "url", "journal"])
        writer.writeheader()
        writer.writerows(rows)


def ensure_page(today: str) -> None:
    subprocess.run([sys.executable, str(DATE_PAGE_SCRIPT), today], cwd=ROOT, check=True)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--date", default=datetime.now(TZ).date().isoformat())
    parser.add_argument("--lookback-days", type=int, default=int(os.getenv("LOOKBACK_DAYS", "120")))
    parser.add_argument("--max-papers", type=int, default=int(os.getenv("MAX_PAPERS", "50")))
    parser.add_argument("--dry-run", action="store_true")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    papers = collect_candidates(args.lookback_days, args.max_papers)
    if not papers:
        if args.dry_run:
            print("No new papers found after de-duplication.")
            return 0
        if issue_exists(args.date):
            ensure_page(args.date)
            print(f"No new papers found, and {args.date} already has a paper-push card.")
            return 0
        previous_date = latest_previous_issue_date(args.date)
        replace_or_prepend_issue(args.date, no_update_issue_block(args.date, previous_date))
        ensure_page(args.date)
        print(
            f"No new papers found after de-duplication. "
            f"Added a no-update card for {args.date}"
            + (f" referencing {previous_date}." if previous_date else ".")
        )
        return 0
    print(f"Found {len(papers)} papers for {args.date}.")
    if args.dry_run:
        for paper in papers[:10]:
            print(f"- {paper.published_month} | {paper.journal} | {paper.title} | {paper.doi}")
        return 0
    replace_or_prepend_issue(args.date, issue_block(args.date, papers))
    write_docx(args.date, papers)
    update_history(args.date, papers)
    ensure_page(args.date)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
