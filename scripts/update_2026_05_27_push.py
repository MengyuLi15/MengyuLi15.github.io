from __future__ import annotations

import csv
import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATE = "2026-05-27"
DOCX = ROOT / "files" / "paper-push" / f"daily_paper_push_{DATE}.docx"
FIGURE = ROOT / "images" / "paper-push" / f"mechanism_{DATE}.png"
HISTORY = ROOT / "推送历史_论文索引.csv"


papers = [
    {
        "group": "Nature 系列",
        "title": "Biochemical remodelling of phytoplankton cell composition under climate change",
        "authors": "Sharoni, S.; Inomura, K.; Dutkiewicz, S.; et al.",
        "journal": "Nature Climate Change",
        "published": "2026-03-31",
        "doi": "10.1038/s41558-026-02598-w",
        "url": "https://www.nature.com/articles/s41558-026-02598-w",
        "tags": "phytoplankton; stoichiometry; carbon pump",
        "summary": "用细胞资源分配模型评估变暖下浮游植物蛋白、糖类和脂类组成变化。结论是高纬和极区蛋白比例可能下降、有机质质量改变；这会影响食物网营养质量和用固定化学计量估算碳输出的可靠性。",
    },
    {
        "group": "Nature 系列",
        "title": "BGC-Argo float reveals shifts in nitrogen-carbon cycling in an oxygen-deficient zone",
        "authors": "Bif, M. B.; Kelly, C.; Altabet, M. A.; et al.",
        "journal": "Communications Earth & Environment",
        "published": "2026-04-06",
        "doi": "10.1038/s43247-026-03410-5",
        "url": "https://www.nature.com/articles/s43247-026-03410-5",
        "tags": "BGC-Argo; microbial carbon; nitrogen cycling; ODZ",
        "summary": "用一支 BGC-Argo 浮标在东热带北太平洋低氧区做近三年剖面观测，联合氧、硝酸盐、亚硝酸盐、pH、叶绿素和 POC。结果显示 BGC-Argo 可用于诊断氮氧化还原过程与碳循环耦合，而不只是看 Chl-a 或 POC 代理量。",
    },
    {
        "group": "Nature 系列",
        "title": "Global declines in net primary production in the ocean color era",
        "authors": "Silsbe, G. M.; Fox, J.; Westberry, T. K.; et al.",
        "journal": "Nature Communications",
        "published": "2025-07-01",
        "doi": "10.1038/s41467-025-60906-y",
        "url": "https://www.nature.com/articles/s41467-025-60906-y",
        "tags": "ocean colour; NPP; climate trend",
        "summary": "用 SeaWiFS 和 MODIS-Aqua 的 25 年海色记录重估全球 NPP 变化。主要结论是近一半海域 NPP 显著下降，集中在热带和副热带层化海区，机制指向营养盐限制增强，是海色监测 carbon pump 生产端变化的关键论文。",
    },
    {
        "group": "Nature 系列",
        "title": "Marine heatwaves modulate food webs and carbon transport processes",
        "authors": "Bif, M. B.; Kellogg, C. T. E.; Huang, Y.; et al.",
        "journal": "Nature Communications",
        "published": "2025-10-06",
        "doi": "10.1038/s41467-025-63605-w",
        "url": "https://www.nature.com/articles/s41467-025-63605-w",
        "tags": "BGC-Argo; marine heatwaves; POC; microbial ecology",
        "summary": "把 Line P 附近 BGC-Argo、色素和 DNA metabarcoding 结合分析。结果显示 2015 和 2019 海洋热浪期间 POC 异常升高，但小颗粒在中层累积并缓慢再矿化，未必提高深海封存，反而可能降低碳泵效率。",
    },
    {
        "group": "Nature 系列",
        "title": "Marine heatwaves are shaping the vertical structure of phytoplankton in the global ocean",
        "authors": "Ma, X.; Chen, G.",
        "journal": "Communications Earth & Environment",
        "published": "2025-08-29",
        "doi": "10.1038/s43247-025-02718-y",
        "url": "https://www.nature.com/articles/s43247-025-02718-y",
        "tags": "BGC-Argo; marine heatwaves; vertical chlorophyll; DCM",
        "summary": "利用 2008-2024 年全球 BGC-Argo 剖面研究 MHW 对浮游植物垂向结构的影响。结论是响应常在亚表层强于表层，可分为增强、减弱和亚表层反转等类型，深叶绿素极大层会被显著重塑。",
    },
    {
        "group": "Nature 系列",
        "title": "The oceanic physical injection pump of organic carbon",
        "authors": "Bellacicco, M.; Marullo, S.; Dall'Olmo, G.; et al.",
        "journal": "Nature Communications",
        "published": "2025-08-02",
        "doi": "10.1038/s41467-025-62363-z",
        "url": "https://www.nature.com/articles/s41467-025-62363-z",
        "tags": "carbon pump; POC/DOC; physical injection pump",
        "summary": "单独量化 biological carbon pump 之外的物理注入泵。结果表明混合、夹卷和输运造成的 POC/DOC 输出不可忽略，carbon pump 研究不能只关注重力沉降。",
    },
    {
        "group": "Science 系列",
        "title": "Antarctic krill vertical migrations modulate seasonal carbon export",
        "authors": "Smith, A. J. R.; Wotherspoon, S.; Ratnarajah, L.; et al.",
        "journal": "Science",
        "published": "2025",
        "doi": "10.1126/science.adq5564",
        "url": "https://www.science.org/doi/10.1126/science.adq5564",
        "tags": "carbon export; vertical migration; Southern Ocean",
        "summary": "研究南大洋磷虾垂向迁移对季节性碳输出的调制。它不是 BGC-Argo 论文，但补上了 migrant pump 这条碳输出通道，对解释 Southern Ocean carbon export 很重要。",
    },
    {
        "group": "Science 系列",
        "title": "Declining ocean greenness and phytoplankton blooms in low to mid-latitudes under a warming climate",
        "authors": "Nature Index records 8 authors",
        "journal": "Science Advances",
        "published": "2025-10-17",
        "doi": "10.1126/sciadv.adx4857",
        "url": "https://www.science.org/doi/10.1126/sciadv.adx4857",
        "tags": "ocean colour; phytoplankton blooms; warming",
        "summary": "用海色长期记录分析中低纬海洋绿度和藻华变化。结果指向变暖背景下中低纬表层绿度和藻华减弱，可作为 ocean colour 趋势监测和浮游植物生产变化的重点跟踪。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Decoupled timescales of organic carbon and phosphorus recycling in the surface ocean",
        "authors": "PNAS metadata",
        "journal": "Proceedings of the National Academy of Sciences",
        "published": "2026",
        "doi": "10.1073/pnas.2514991123",
        "url": "https://www.pnas.org/doi/10.1073/pnas.2514991123",
        "tags": "organic carbon; phosphorus recycling; microbial carbon",
        "summary": "聚焦表层海洋有机碳和磷再循环时间尺度的解耦。结论提示微生物再矿化和元素循环不能简单按固定比例处理，对 microbial carbon 和 carbon pump 机制解释很有价值。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Global warming drives more intense, widespread, and frequent marine heatwaves",
        "authors": "X. Li; C. Zhang; M. Newman; F.-F. Jin",
        "journal": "Proceedings of the National Academy of Sciences",
        "published": "2025",
        "doi": "10.1073/pnas.2413505122",
        "url": "https://www.pnas.org/doi/10.1073/pnas.2413505122",
        "tags": "marine heatwaves; warming; global trend",
        "summary": "定量说明全球变暖会推动海洋热浪更强、更广、更频繁。它是 MHW 生态影响研究的物理背景论文，可帮助解释为什么垂向浮游植物响应会越来越重要。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Characterization of phytoplankton-excreted metabolites using high-resolution mass spectrometry and stable isotope tracing",
        "authors": "PNAS metadata",
        "journal": "Proceedings of the National Academy of Sciences",
        "published": "2026",
        "doi": "10.1073/pnas.2531765123",
        "url": "https://www.pnas.org/doi/10.1073/pnas.2531765123",
        "tags": "phytoplankton metabolites; dissolved organic matter; microbial loop",
        "summary": "用高分辨率质谱和稳定同位素追踪浮游植物释放代谢物。结论可帮助理解初级生产如何进入 DOM 和微生物环，对 microbial carbon 和碳泵非颗粒通道很相关。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "PACE/OCI retrievals of the spectral particulate backscattering coefficient in the global ocean using mixture density networks",
        "authors": "Remote Sensing of Environment metadata",
        "journal": "Remote Sensing of Environment",
        "published": "2026",
        "doi": "10.1016/j.rse.2026.115327",
        "url": "https://doi.org/10.1016/j.rse.2026.115327",
        "tags": "PACE; ocean colour; backscattering; machine learning",
        "summary": "面向 PACE/OCI 反演全球海洋颗粒后向散射光谱，用 mixture density networks 表征不确定性。对 BGC-Argo bbp 交叉验证、bio-optic 参数化和 POC 遥感很重要。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Global surface phytoplankton community structure from ocean color radiometry",
        "authors": "Remote Sensing of Environment metadata",
        "journal": "Remote Sensing of Environment",
        "published": "2025",
        "doi": "10.1016/j.rse.2025.115174",
        "url": "https://doi.org/10.1016/j.rse.2025.115174",
        "tags": "ocean colour; phytoplankton community structure; remote sensing",
        "summary": "用海色辐射计估计全球表层浮游植物群落结构。它把传统 Chl-a 监测推进到群落组成层面，适合与 BGC-Argo 垂向剖面结合，分析表层与 DCM 群落差异。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "SeaHawk CubeSat captures ocean color with high spatial and temporal resolution",
        "authors": "Remote Sensing of Environment metadata",
        "journal": "Remote Sensing of Environment",
        "published": "2025",
        "doi": "10.1016/j.rse.2025.115111",
        "url": "https://doi.org/10.1016/j.rse.2025.115111",
        "tags": "ocean colour; high-resolution observation; CubeSat",
        "summary": "展示 SeaHawk 小卫星在高时空分辨率海色观测中的能力。对近岸和小尺度藻华、锋面、热浪响应监测有用，可补足传统海色产品空间分辨率不足的问题。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Phytoplankton flexible pigment content complicates modeling of primary production and carbon fluxes",
        "authors": "Global Change Biology metadata",
        "journal": "Global Change Biology",
        "published": "2026",
        "doi": "10.1111/gcb.70671",
        "url": "https://doi.org/10.1111/gcb.70671",
        "tags": "phytoplankton pigments; primary production; carbon flux",
        "summary": "讨论浮游植物色素含量可塑性如何影响初级生产和碳通量建模。结论提醒 Chl-a 到碳或 NPP 的转换会受 photoacclimation 影响，特别适合和 BGC-Argo 垂向 Chl/光场联用。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Modeling Organic Carbon Flux",
        "authors": "Annual Review of Marine Science metadata",
        "journal": "Annual Review of Marine Science",
        "published": "2026",
        "doi": "10.1146/annurev-marine-022123-102516",
        "url": "https://doi.org/10.1146/annurev-marine-022123-102516",
        "tags": "carbon flux; review; model synthesis",
        "summary": "综述有机碳通量建模的主要框架、观测约束和不确定性。它适合作为 carbon pump 研究的路线图，尤其能帮助判断 BGC-Argo、沉积物捕集器和卫星产品各自约束的是哪一段通量。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Two Decades of Increase in Southern Ocean Net Community Production Revealed by BGC-Argo Floats",
        "authors": "Liniger, G.; Sharp, J. D.; Takeshita, Y.; Johnson, K. S.",
        "journal": "Global Biogeochemical Cycles",
        "published": "2025",
        "doi": "10.1029/2024GB008371",
        "url": "https://repository.library.noaa.gov/view/noaa/71386",
        "tags": "BGC-Argo; NCP; Southern Ocean",
        "summary": "用 BGC-Argo 约束神经网络重构南大洋硝酸盐，再由季节性消耗估算 ANCP。结论是 2004-2022 年南大洋 ANCP 约 3.91 PgC/yr 且呈增长趋势，传统方法可能低估深层碳输出。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Southern Ocean Carbon Export Revealed by Backscatter and Oxygen Measurements From BGC-Argo Floats",
        "authors": "Liniger, G.; Moreau, S.; Lannuzel, D.; Carranza, M. M.; Strutton, P. G.",
        "journal": "Global Biogeochemical Cycles",
        "published": "2025",
        "doi": "10.1029/2024GB008193",
        "url": "https://doi.org/10.1029/2024GB008193",
        "tags": "BGC-Argo; backscatter; oxygen; carbon export",
        "summary": "用 212 个 BGC-Argo 浮标的 backscatter 和氧变化估算南大洋碳输出。结果给出约 2.69 Pg C/yr 的下沉输出，并指出海冰区贡献约 8%，对南大洋碳泵区域差异很有参考价值。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Deep Learning Estimation of Global Surface Chlorophyll a Concentration From Satellite Remote Sensing",
        "authors": "Geophysical Research Letters metadata",
        "journal": "Geophysical Research Letters",
        "published": "2026",
        "doi": "10.1029/2025GL120669",
        "url": "https://doi.org/10.1029/2025GL120669",
        "tags": "ocean colour; deep learning; chlorophyll",
        "summary": "用深度学习改进卫星表层 Chl-a 估算。它对海色产品算法发展很直接，后续可与 BGC-Argo 剖面比较，评估表层 Chl-a 误差如何传播到 NPP 和碳输出估算。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Seasonal Carbon Uptake by Phytoplankton in Coastal Waters of the Western Antarctic Peninsula",
        "authors": "Geophysical Research Letters metadata",
        "journal": "Geophysical Research Letters",
        "published": "2025",
        "doi": "10.1029/2024GL112446",
        "url": "https://doi.org/10.1029/2024GL112446",
        "tags": "phytoplankton; carbon uptake; Antarctic Peninsula",
        "summary": "研究西南极半岛沿岸浮游植物季节性碳吸收。它有助于解释高纬边缘海的生产力季节窗口，也能与南大洋 BGC-Argo NCP 和 carbon export 结果相互印证。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Unprecedented phytoplankton bloom induced by the earliest and strongest marine heatwave in the Ross Sea",
        "authors": "Geophysical Research Letters metadata",
        "journal": "Geophysical Research Letters",
        "published": "2025",
        "doi": "10.1029/2024GL111264",
        "url": "https://doi.org/10.1029/2024GL111264",
        "tags": "marine heatwave; phytoplankton bloom; Ross Sea",
        "summary": "报道罗斯海早发且强烈的海洋热浪诱发异常浮游植物藻华。它说明 MHW 不总是抑制浮游植物，在高纬海区也可能通过海冰、层化和光照窗口改变 bloom timing 与强度。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Subsurface Chlorophyll Maxima Reduce the Performance of Non-Photochemical Quenching Corrections",
        "authors": "Journal of Geophysical Research: Oceans metadata",
        "journal": "Journal of Geophysical Research: Oceans",
        "published": "2026",
        "doi": "10.1029/2025JC023249",
        "url": "https://doi.org/10.1029/2025JC023249",
        "tags": "BGC-Argo; chlorophyll fluorescence; NPQ; DCM",
        "summary": "指出亚表层叶绿素极大层会削弱常用非光化学淬灭校正的表现。对 BGC-Argo Chl fluorescence 数据处理非常关键，尤其是研究 DCM 和热浪垂向响应时需要注意。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "BGC-Argo+: Tiered standard extensions for the global biogeochemical-Argo array",
        "authors": "ESSD discussion metadata",
        "journal": "Earth System Science Data Discussions",
        "published": "2026",
        "doi": "10.5194/essd-2026-311",
        "url": "https://doi.org/10.5194/essd-2026-311",
        "tags": "BGC-Argo; data system; sensors; methods",
        "summary": "提出 BGC-Argo+ 分层扩展方案，讨论全球 BGC-Argo 阵列未来传感器和标准化扩展。它偏方法和观测系统设计，但对后续 bio-optic、carbon pump 和微生物过程观测非常重要。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "A global decadal data record of 8-day resolution chlorophyll-a concentration from multi-source satellite measurements",
        "authors": "ESSD metadata",
        "journal": "Earth System Science Data",
        "published": "2026",
        "doi": "10.5194/essd-18-569-2026",
        "url": "https://doi.org/10.5194/essd-18-569-2026",
        "tags": "ocean colour; chlorophyll dataset; multi-sensor",
        "summary": "发布全球 8 日分辨率多源卫星 Chl-a 数据记录。这个数据集适合做长期海色趋势、MHW 前后变化和 BGC-Argo 剖面匹配，是日常论文监测中的数据源型重点。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "The DYPHYRAD dataset: carbon, nitrogen, phosphorus, and phytoplankton groups from the Bay of Brest",
        "authors": "ESSD metadata",
        "journal": "Earth System Science Data",
        "published": "2026",
        "doi": "10.5194/essd-18-1877-2026",
        "url": "https://doi.org/10.5194/essd-18-1877-2026",
        "tags": "phytoplankton groups; nutrients; carbon; dataset",
        "summary": "发布布雷斯特湾长期碳、氮、磷和浮游植物类群数据集。虽是区域数据，但可用于验证浮游植物群落变化、营养盐控制和海色类群算法。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Southern Ocean particulate organic matter fluxes estimated from in situ optical observations",
        "authors": "Biogeosciences metadata",
        "journal": "Biogeosciences",
        "published": "2026",
        "doi": "10.5194/bg-23-2179-2026",
        "url": "https://doi.org/10.5194/bg-23-2179-2026",
        "tags": "Southern Ocean; optical observations; POM flux",
        "summary": "用原位光学观测估算南大洋颗粒有机物通量。它直接连接 bio-optic proxy 和 carbon export，是 BGC-Argo bbp/POC 推断输出通量的重要参考。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "A first-order description of the relationship between the chlorophyll maximum layer and colored dissolved organic matter",
        "authors": "Biogeosciences metadata",
        "journal": "Biogeosciences",
        "published": "2026",
        "doi": "10.5194/bg-23-3073-2026",
        "url": "https://doi.org/10.5194/bg-23-3073-2026",
        "tags": "chlorophyll maximum; CDOM; bio-optics",
        "summary": "建立叶绿素最大层和 CDOM 之间关系的一阶描述。它对解释 DCM 附近吸收光学信号很有用，也能帮助区分 Chl-a、CDOM 和颗粒吸收对 ocean colour/bio-optic 反演的贡献。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Phytoplankton functional groups in the recurrent Kerguelen bloom respond differently to iron supply",
        "authors": "Biogeosciences metadata",
        "journal": "Biogeosciences",
        "published": "2026",
        "doi": "10.5194/bg-23-639-2026",
        "url": "https://doi.org/10.5194/bg-23-639-2026",
        "tags": "phytoplankton functional groups; iron; Southern Ocean bloom",
        "summary": "研究 Kerguelen recurrent bloom 中不同浮游植物功能群对铁供应的差异响应。对南大洋生产力、群落组成和碳输出效率之间的机制联系很有参考价值。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Marine heatwaves impact on phytoplankton communities in Drake Passage",
        "authors": "Biogeosciences metadata",
        "journal": "Biogeosciences",
        "published": "2025",
        "doi": "10.5194/bg-22-7205-2025",
        "url": "https://doi.org/10.5194/bg-22-7205-2025",
        "tags": "marine heatwaves; phytoplankton communities; Southern Ocean",
        "summary": "分析 Drake Passage 海洋热浪对浮游植物群落的影响。它补充了 BGC-Argo 热浪垂向结构研究中的群落组成视角，尤其适合关注南大洋高纬 MHW 生态效应。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Lifting the lid on marine heatwaves: Global patterns of subsurface ocean warming events",
        "authors": "Progress in Oceanography metadata",
        "journal": "Progress in Oceanography",
        "published": "2025",
        "doi": "10.1016/j.pocean.2025.103539",
        "url": "https://doi.org/10.1016/j.pocean.2025.103539",
        "tags": "marine heatwaves; subsurface warming; vertical structure",
        "summary": "从全球角度揭示亚表层海洋增暖事件的空间格局。它是理解 MHW 垂向生态影响的物理底座，能帮助解释为什么只看 SST 会漏掉 DCM 和亚表层 Chl 响应。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Climate variability affects phytoplankton phenology on the South West Atlantic shelf",
        "authors": "Progress in Oceanography metadata",
        "journal": "Progress in Oceanography",
        "published": "2025",
        "doi": "10.1016/j.pocean.2025.103583",
        "url": "https://doi.org/10.1016/j.pocean.2025.103583",
        "tags": "phytoplankton phenology; climate variability; shelf ocean",
        "summary": "研究气候变率如何影响西南大西洋陆架浮游植物物候。对做海色时间序列、bloom timing 和气候模态影响分析很有用。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Bio-optical properties in summer surface waters of the southern Chukchi Sea and northern Bering Sea",
        "authors": "Deep Sea Research Part I metadata",
        "journal": "Deep Sea Research Part I: Oceanographic Research Papers",
        "published": "2025",
        "doi": "10.1016/j.dsr.2025.104458",
        "url": "https://doi.org/10.1016/j.dsr.2025.104458",
        "tags": "bio-optics; Arctic; ocean colour validation",
        "summary": "研究楚科奇海南部和白令海北部夏季表层水体 bio-optical properties。它可用于高纬海色算法验证，并帮助解释 CDOM、颗粒和 Chl-a 对遥感反射率的影响。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Top-of-atmosphere hyperspectral ocean color estimation for NOAA-21 VIIRS",
        "authors": "IEEE TGRS metadata",
        "journal": "IEEE Transactions on Geoscience and Remote Sensing",
        "published": "2026",
        "doi": "10.1109/TGRS.2026.3670718",
        "url": "https://doi.org/10.1109/TGRS.2026.3670718",
        "tags": "ocean colour; VIIRS; hyperspectral estimation",
        "summary": "从 NOAA-21 VIIRS 顶层大气信号估计高光谱海色信息。它偏算法和传感器应用，对未来多传感器海色融合、PACE/VIIRS 交叉应用有参考价值。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "A multi-task learning method for ocean color remote sensing products",
        "authors": "IEEE TGRS metadata",
        "journal": "IEEE Transactions on Geoscience and Remote Sensing",
        "published": "2026",
        "doi": "10.1109/TGRS.2026.3675239",
        "url": "https://doi.org/10.1109/TGRS.2026.3675239",
        "tags": "ocean colour; multi-task learning; remote sensing products",
        "summary": "提出多任务学习方法联合生成海色遥感产品。对 Chl-a、IOPs、POC 等多变量联合反演有启发，适合和 BGC-Argo 多参数剖面做一致性检验。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "A Guide to Field Campaign Design for Satellite Calibration and Validation",
        "authors": "Journal of Atmospheric and Oceanic Technology metadata",
        "journal": "Journal of Atmospheric and Oceanic Technology",
        "published": "2026",
        "doi": "10.1175/JTECH-D-25-0079.1",
        "url": "https://doi.org/10.1175/JTECH-D-25-0079.1",
        "tags": "satellite validation; field campaign; PACE",
        "summary": "讨论卫星校准验证野外航次/观测设计。它对 ocean colour 和 BGC-Argo 匹配验证很实用，尤其适合规划 PACE、BioGeoChemical Argo 与船基光学观测协同。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Marine primary production and carbon export in a western Mediterranean oligotrophic system using sediment traps",
        "authors": "Frontiers in Marine Science metadata",
        "journal": "Frontiers in Marine Science",
        "published": "2026",
        "doi": "10.3389/fmars.2026.1806793",
        "url": "https://doi.org/10.3389/fmars.2026.1806793",
        "tags": "primary production; carbon export; sediment traps",
        "summary": "用沉积物捕集器研究西地中海贫营养系统初级生产和碳输出。它能为 BGC-Argo/光学估算的 carbon export 提供独立对照。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "A neural network emulator for operational marine biogeochemical forecasting in the Black Sea",
        "authors": "Frontiers in Marine Science metadata",
        "journal": "Frontiers in Marine Science",
        "published": "2026",
        "doi": "10.3389/fmars.2026.1760162",
        "url": "https://doi.org/10.3389/fmars.2026.1760162",
        "tags": "biogeochemical forecasting; neural network; operational oceanography",
        "summary": "构建黑海业务化海洋生物地球化学预报的神经网络 emulator。虽然区域是黑海，但方法上与多源观测、BGC-Argo/卫星融合和快速 biogeochemical prediction 有联系。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Biogeochemical response of oceanic subsurface chlorophyll maximum to Hurricane Idalia",
        "authors": "Frontiers in Marine Science metadata",
        "journal": "Frontiers in Marine Science",
        "published": "2026",
        "doi": "10.3389/fmars.2026.1740354",
        "url": "https://doi.org/10.3389/fmars.2026.1740354",
        "tags": "subsurface chlorophyll maximum; storm disturbance; vertical response",
        "summary": "研究飓风 Idalia 对海洋亚表层叶绿素最大层的生物地球化学响应。虽然扰动类型不是 MHW，但同样强调垂向结构和 SCM/DCM，对热浪垂向响应研究有类比价值。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Spectral Dependence of the Depolarization Ratio of Pure Water at Ocean Lidar Wavelengths",
        "authors": "Applied Optics metadata",
        "journal": "Applied Optics",
        "published": "2026",
        "doi": "10.1364/AO.590152",
        "url": "https://doi.org/10.1364/AO.590152",
        "tags": "ocean lidar; pure water optics; bio-optic methods",
        "summary": "研究海洋激光雷达波段纯水退偏比的光谱依赖。它不是生态论文，但对未来三维 ocean colour、lidar 和 BGC-Argo 垂向光学融合有方法学意义。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Spectral variability in particulate light backscattering in the sea",
        "authors": "Optics Express metadata",
        "journal": "Optics Express",
        "published": "2025",
        "doi": "10.1364/OE.579542",
        "url": "https://doi.org/10.1364/OE.579542",
        "tags": "particulate backscattering; bio-optics; ocean colour",
        "summary": "研究海洋颗粒后向散射的光谱变化。它直接关系到 bbp、POC 和颗粒群落性质的遥感/剖面反演，是 bio-optic 和 BGC-Argo 光学解释的基础论文。",
    },
    {
        "group": "重点期刊：按影响力和相关性排序",
        "title": "Advancing Earth System Science With the NASA Plankton, Aerosol, Cloud, Ocean Ecosystem (PACE) Satellite Mission",
        "authors": "Oceanography metadata",
        "journal": "Oceanography",
        "published": "2026",
        "doi": "10.5670/oceanog.2026.e204",
        "url": "https://doi.org/10.5670/oceanog.2026.e204",
        "tags": "PACE; ocean colour; hyperspectral remote sensing",
        "summary": "介绍 PACE 任务如何推动地球系统科学，尤其是浮游植物、气溶胶、云和海洋生态系统观测。它是 PACE 背景和应用框架论文，对 ocean colour 与 BGC-Argo 融合很重要。",
    },
]


def q(value: str) -> str:
    return json.dumps(value, ensure_ascii=False)


def write_yaml() -> None:
    lines = [
        f"- date: {q(DATE)}",
        f"  title: {q('每日论文推送：BGC-Argo、海色、海洋热浪与碳泵')}",
        f"  summary: {q(f'本期按更新后的规则重跑：Nature/Science 系列优先，其余重点期刊按影响力与主题相关性排序；历史去重后保留 {len(papers)} 篇，不超过每日 50 篇上限。')}",
        f"  docx: {q('/files/paper-push/' + DOCX.name)}",
        f"  figure: {q('/images/paper-push/' + FIGURE.name)}",
        "  papers:",
    ]
    for paper in papers:
        lines.append(f"    - title: {q(paper['title'])}")
        for key in ["authors", "journal", "published", "doi", "url", "group", "tags", "summary"]:
            lines.append(f"      {key}: {q(paper[key])}")
    (ROOT / "_data" / "paper_pushes.yml").write_text("\n".join(lines) + "\n", encoding="utf-8")


def set_font(run, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def write_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(31, 77, 120)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 11.5, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("每日论文推送")
    set_font(r)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(23, 50, 77)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"检索日期：{DATE} | 主题：BGC-Argo / ocean colour / marine heatwaves / phytoplankton / carbon pump")
    set_font(r)
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(85, 105, 125)

    doc.add_heading("今日总览", level=1)
    p = doc.add_paragraph()
    r = p.add_run(f"本期按更新后的重点期刊列表重跑。Nature/Science 系列优先，其余期刊按影响力和主题相关性排序；历史去重后保留 {len(papers)} 篇，未超过每日 50 篇上限。重点信号包括：MHW 研究正在转向亚表层与 DCM，海色研究正向 PACE/高光谱/机器学习推进，BGC-Argo 光学与氧/硝酸盐组合正在更直接约束 NCP、POC export 和微生物过程。")
    set_font(r)

    if FIGURE.exists():
        doc.add_picture(str(FIGURE), width=Inches(6.4))
        cap = doc.add_paragraph("图 1. 自绘机制示意图：海色表层观测、BGC-Argo 垂向剖面、热浪扰动与碳泵路径的关系。")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            set_font(run)
            run.font.size = Pt(9)

    doc.add_heading("论文速读", level=1)
    current_group = None
    for i, paper in enumerate(papers, 1):
        if paper["group"] != current_group:
            current_group = paper["group"]
            doc.add_heading(current_group, level=2)
        p = doc.add_paragraph(style="Heading 3")
        r = p.add_run(f"{i}. {paper['title']}")
        set_font(r)
        r.bold = True
        p = doc.add_paragraph()
        r = p.add_run(f"{paper['authors']} | {paper['journal']} | {paper['published']} | DOI: {paper['doi']}")
        set_font(r)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(85, 105, 125)
        p = doc.add_paragraph()
        r = p.add_run(f"关键词：{paper['tags']}")
        set_font(r)
        r.bold = True
        r.font.size = Pt(9.5)
        p = doc.add_paragraph()
        r = p.add_run(paper["summary"])
        set_font(r)
        p = doc.add_paragraph()
        r = p.add_run(f"链接：{paper['url']}")
        set_font(r)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(5, 99, 193)

    doc.add_heading("检索与排序说明", level=1)
    p = doc.add_paragraph()
    r = p.add_run("检索覆盖 Google Scholar 可见结果、期刊官网、出版社 DOI 页、Crossref/AGU/Copernicus/ScienceDirect/Nature/PNAS 等公开页面。排序为 Nature 系列、Science 系列、其余重点期刊按影响因子/影响力与主题相关性近似排序。")
    set_font(r)
    doc.save(DOCX)


def write_history() -> None:
    rows = []
    if HISTORY.exists():
        with HISTORY.open("r", encoding="utf-8-sig", newline="") as fh:
            rows = list(csv.DictReader(fh))
    seen = {(row.get("doi") or "").lower() for row in rows}
    for paper in papers:
        if paper["doi"].lower() not in seen:
            rows.append({
                "date": DATE,
                "title": paper["title"],
                "doi": paper["doi"],
                "url": paper["url"],
                "journal": paper["journal"],
            })
            seen.add(paper["doi"].lower())
    with HISTORY.open("w", encoding="utf-8-sig", newline="") as fh:
        writer = csv.DictWriter(fh, fieldnames=["date", "title", "doi", "url", "journal"])
        writer.writeheader()
        writer.writerows(rows)


def main() -> None:
    DOCX.parent.mkdir(parents=True, exist_ok=True)
    FIGURE.parent.mkdir(parents=True, exist_ok=True)
    old_figure = ROOT / "images" / "paper-push" / "机制示意图_2026-05-27.png"
    if old_figure.exists() and not FIGURE.exists():
        shutil.copyfile(old_figure, FIGURE)
    write_yaml()
    write_docx()
    write_history()
    print(f"Wrote {len(papers)} papers")
    print(DOCX)


if __name__ == "__main__":
    main()
